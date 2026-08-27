import re
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.image.exceptions import UnrecognizedImageError
from docx.image.image import Image as DocxImage
from docx.shared import Pt

PLACEHOLDER_PATTERN = re.compile(r'\{\{\s*([a-zA-Z0-9_]+)\s*\}\}')


def _iter_paragraphs(parent):
    """Yield every paragraph in parent (Document, _Cell, _Header or _Footer),
    descending recursively into nested tables."""
    for p in parent.paragraphs:
        yield p
    tables = getattr(parent, 'tables', None)
    if tables:
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from _iter_paragraphs(cell)


def _iter_all_paragraphs(doc):
    yield from _iter_paragraphs(doc)
    for section in doc.sections:
        yield from _iter_paragraphs(section.header)
        yield from _iter_paragraphs(section.footer)


def replace_placeholders_in_paragraph(paragraph, context):
    """Replace every {{key}} found in paragraph with context[key], handling
    placeholders split across multiple runs by Word. Returns the set of
    placeholder keys actually replaced (present in context)."""
    used = set()
    while True:
        runs = paragraph.runs
        if not runs:
            break
        texts = [r.text for r in runs]
        full_text = ''.join(texts)
        match = PLACEHOLDER_PATTERN.search(full_text)
        if not match:
            break
        key = match.group(1)
        start, end = match.start(), match.end()

        offsets = []
        pos = 0
        for t in texts:
            offsets.append((pos, pos + len(t)))
            pos += len(t)

        covered = [i for i, (s, e) in enumerate(offsets) if e > start and s < end]
        if not covered:
            break
        first_idx, last_idx = covered[0], covered[-1]

        prefix = texts[first_idx][:start - offsets[first_idx][0]]
        suffix = texts[last_idx][end - offsets[last_idx][0]:]

        if key in context:
            value = str(context[key])
            used.add(key)
        else:
            value = match.group(0)  # leave unresolved placeholder untouched

        if first_idx == last_idx:
            runs[first_idx].text = prefix + value + suffix
        else:
            runs[first_idx].text = prefix + value
            for i in covered[1:-1]:
                runs[i].text = ''
            runs[last_idx].text = suffix

        if key not in context:
            # avoid infinite loop when a placeholder has no value: stop trying
            # to re-match the same unresolved text by breaking out here.
            break
    return used


def replace_placeholders_in_doc(doc, context):
    used = set()
    for paragraph in _iter_all_paragraphs(doc):
        used |= replace_placeholders_in_paragraph(paragraph, context)
    return used


def discover_placeholders_in_docx(path):
    """Scan a .docx template and return the ordered list of distinct
    placeholder keys found (first-appearance order)."""
    doc = Document(path)
    seen = []
    seen_set = set()
    for paragraph in _iter_all_paragraphs(doc):
        full_text = ''.join(r.text for r in paragraph.runs)
        for m in PLACEHOLDER_PATTERN.finditer(full_text):
            key = m.group(1)
            if key not in seen_set:
                seen_set.add(key)
                seen.append(key)
    return seen


def validate_image_file(path):
    """Verifica che path sia un'immagine in un formato che python-docx sa
    incorporare (PNG/JPEG/BMP/GIF/TIFF), senza dipendere da Pillow."""
    try:
        DocxImage.from_file(path)
    except UnrecognizedImageError as exc:
        raise ValueError(
            'Il file scelto non è un formato immagine supportato (usa PNG, JPG, BMP, GIF o TIFF).'
        ) from exc
    except (FileNotFoundError, OSError) as exc:
        raise ValueError(f'Impossibile leggere il file immagine: {exc}') from exc


def _clear_header_content(header):
    """Rimuove ogni paragrafo/tabella preesistente nell'header e lascia un
    unico paragrafo vuoto pronto per ricevere l'immagine di intestazione.
    Un <w:hdr> deve sempre contenere almeno un <w:p>."""
    hdr_element = header._element
    for child in list(hdr_element):
        if child.tag.endswith('}p') or child.tag.endswith('}tbl'):
            hdr_element.remove(child)
    return header.add_paragraph()


def apply_letterhead(doc, image_path):
    """Sostituisce l'header di ogni sezione con l'immagine data, centrata,
    larga quanto l'area stampabile (larghezza pagina meno margini) di quella
    sezione. Non tocca footer né corpo del documento."""
    for section in doc.sections:
        section.header.is_linked_to_previous = False
        paragraph = _clear_header_content(section.header)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        available_width = section.page_width - section.left_margin - section.right_margin
        run = paragraph.add_run()
        run.add_picture(image_path, width=available_width)


def apply_typography(doc, font_family=None, font_size_pt=None, line_spacing=None):
    """Applica font/dimensione/interlinea solo al CORPO del documento (mai
    header/footer). Tocca solo run.font.name/size, mai bold/italic/underline,
    così l'enfasi già presente nel template resta intatta."""
    for paragraph in _iter_paragraphs(doc):
        if line_spacing is not None:
            paragraph.paragraph_format.line_spacing = line_spacing
        for run in paragraph.runs:
            if font_family:
                run.font.name = font_family
            if font_size_pt:
                run.font.size = Pt(font_size_pt)


def generate_document(template_path, output_path, context, letterhead_path=None,
                       font_family=None, font_size_pt=None, line_spacing=None):
    doc = Document(template_path)
    used = replace_placeholders_in_doc(doc, context)
    if letterhead_path:
        apply_letterhead(doc, letterhead_path)
    if font_family or font_size_pt or line_spacing:
        apply_typography(doc, font_family, font_size_pt, line_spacing)
    doc.save(output_path)
    return used


def format_value(value, tipo_campo):
    if value is None or value == '':
        return ''
    if tipo_campo == 'data':
        try:
            return datetime.strptime(str(value), '%Y-%m-%d').strftime('%d/%m/%Y')
        except ValueError:
            return str(value)
    if tipo_campo == 'importo':
        try:
            cents = int(value)
        except (TypeError, ValueError):
            return str(value)
        euros = cents / 100
        formatted = f'{euros:,.2f}'.replace(',', '#').replace('.', ',').replace('#', '.')
        return f'{formatted} €'
    if tipo_campo == 'numero':
        try:
            f = float(value)
        except (TypeError, ValueError):
            return str(value)
        if f == int(f):
            return str(int(f))
        return str(f).replace('.', ',')
    return str(value)
