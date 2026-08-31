"""Genera un documento .docx da un template + i dati di una pratica.

Funzione serverless Python (Vercel Python Runtime). Tutto in UN SOLO file:
Vercel non include automaticamente altri file .py della stessa cartella nel
pacchetto di una funzione, quindi qualunque `from _modulo import x` fallisce
con ModuleNotFoundError a runtime (verificato con un test diagnostico
dedicato) — niente import locali, mai.

Verifica sempre l'identità di chi chiama tramite Supabase Auth (mai fidandosi
di un ID studio passato dal client) e non fa altro che questo più
leggere/scrivere lo Storage cifrato: non ha accesso diretto al database
applicativo oltre alle query dirette a Postgrest necessarie per raccogliere
il contesto del documento.
"""
import base64
import json
import os
import re
import tempfile
import urllib.error
import urllib.request
import uuid
from datetime import datetime
from http.server import BaseHTTPRequestHandler

from cryptography.hazmat.primitives import hashes
from cryptography.hazmat.primitives.ciphers.aead import AESGCM
from cryptography.hazmat.primitives.kdf.hkdf import HKDF
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

SUPABASE_URL = os.environ.get('NEXT_PUBLIC_SUPABASE_URL', '')
ANON_KEY = os.environ.get('NEXT_PUBLIC_SUPABASE_ANON_KEY', '')
SERVICE_ROLE_KEY = os.environ.get('SUPABASE_SERVICE_ROLE_KEY', '')
BUCKET = 'documents'

# ---------- cifratura (stessa di src/lib/crypto/docEncryption.ts) ----------

HKDF_SALT = b'themis-doc-key'
IV_LENGTH = 12
SYSTEM_SCOPE = 'system'


def _master_key():
    return base64.b64decode(os.environ['DOCUMENT_ENCRYPTION_MASTER_KEY'])


def derive_key(scope: str) -> bytes:
    hkdf = HKDF(algorithm=hashes.SHA256(), length=32, salt=HKDF_SALT, info=scope.encode('utf-8'))
    return hkdf.derive(_master_key())


def encrypt_bytes(plaintext: bytes, scope: str) -> bytes:
    key = derive_key(scope)
    iv = os.urandom(IV_LENGTH)
    return iv + AESGCM(key).encrypt(iv, plaintext, None)


def decrypt_bytes(blob: bytes, scope: str) -> bytes:
    key = derive_key(scope)
    iv = blob[:IV_LENGTH]
    return AESGCM(key).decrypt(iv, blob[IV_LENGTH:], None)


# ---------- auth (verifica il token contro Supabase, mai fidarsi del client) ----------

def get_studio_id(access_token: str):
    """Studio per cui lavora chi sta chiamando.

    Non è più "l'id dell'utente", perché un collaboratore ha un id proprio
    diverso da quello dello studio: la risposta la dà Postgres con
    studio_corrente(), la stessa identica funzione che regola l'accesso ai
    dati nel resto dell'applicazione — una sola fonte di verità invece di
    due implementazioni da tenere allineate.

    Chiamata con il token dell'utente e non con la chiave di servizio: se il
    token non è valido Supabase risponde 401 e qui si ritorna None, quindi
    resta vero quanto dichiarato in cima al file (verificare sempre chi
    chiama, senza fidarsi di un id passato dal client).
    """
    if not access_token:
        return None
    req = urllib.request.Request(
        f'{SUPABASE_URL}/rest/v1/rpc/studio_corrente', data=b'{}', method='POST'
    )
    req.add_header('apikey', ANON_KEY)
    req.add_header('Authorization', f'Bearer {access_token}')
    req.add_header('Content-Type', 'application/json')
    try:
        with urllib.request.urlopen(req) as resp:
            return json.loads(resp.read().decode('utf-8') or 'null')
    except urllib.error.HTTPError:
        return None


# ---------- storage / rest (via service role key) ----------

def download_object(storage_path: str) -> bytes:
    req = urllib.request.Request(f'{SUPABASE_URL}/storage/v1/object/{BUCKET}/{storage_path}')
    req.add_header('apikey', SERVICE_ROLE_KEY)
    req.add_header('Authorization', f'Bearer {SERVICE_ROLE_KEY}')
    with urllib.request.urlopen(req) as resp:
        return resp.read()


def upload_object(storage_path: str, data: bytes):
    req = urllib.request.Request(f'{SUPABASE_URL}/storage/v1/object/{BUCKET}/{storage_path}', data=data, method='POST')
    req.add_header('apikey', SERVICE_ROLE_KEY)
    req.add_header('Authorization', f'Bearer {SERVICE_ROLE_KEY}')
    req.add_header('Content-Type', 'application/octet-stream')
    req.add_header('x-upsert', 'true')
    with urllib.request.urlopen(req) as resp:
        resp.read()


def rest_get(path: str):
    req = urllib.request.Request(f'{SUPABASE_URL}{path}')
    req.add_header('apikey', SERVICE_ROLE_KEY)
    req.add_header('Authorization', f'Bearer {SERVICE_ROLE_KEY}')
    with urllib.request.urlopen(req) as resp:
        return json.loads(resp.read().decode('utf-8') or '[]')


def rest_post(path: str, body):
    data = json.dumps(body).encode('utf-8')
    req = urllib.request.Request(f'{SUPABASE_URL}{path}', data=data, method='POST')
    req.add_header('apikey', SERVICE_ROLE_KEY)
    req.add_header('Authorization', f'Bearer {SERVICE_ROLE_KEY}')
    req.add_header('Content-Type', 'application/json')
    req.add_header('Prefer', 'return=representation')
    with urllib.request.urlopen(req) as resp:
        return json.loads(resp.read().decode('utf-8') or '[]')


def _first(rows):
    return rows[0] if rows else None


# ---------- generazione docx (identico a backend/docgen.py) ----------

PLACEHOLDER_PATTERN = re.compile(r'\{\{\s*([a-zA-Z0-9_]+)\s*\}\}')


def _iter_paragraphs(parent):
    for p in parent.paragraphs:
        yield p
    tables = getattr(parent, 'tables', None)
    if tables:
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from _iter_paragraphs(cell)


def _iter_all_paragraphs(doc):
    yield from _iter_paragraphs(doc)
    for section in doc.sections:
        yield from _iter_paragraphs(section.header)
        yield from _iter_paragraphs(section.footer)


def replace_placeholders_in_paragraph(paragraph, context):
    used = set()
    while True:
        runs = paragraph.runs
        if not runs:
            break
        texts = [r.text for r in runs]
        full_text = ''.join(texts)
        match = PLACEHOLDER_PATTERN.search(full_text)
        if not match:
            break
        key = match.group(1)
        start, end = match.start(), match.end()

        offsets = []
        pos = 0
        for t in texts:
            offsets.append((pos, pos + len(t)))
            pos += len(t)

        covered = [i for i, (s, e) in enumerate(offsets) if e > start and s < end]
        if not covered:
            break
        first_idx, last_idx = covered[0], covered[-1]

        prefix = texts[first_idx][:start - offsets[first_idx][0]]
        suffix = texts[last_idx][end - offsets[last_idx][0]:]

        if key in context:
            value = str(context[key])
            used.add(key)
        else:
            value = match.group(0)

        if first_idx == last_idx:
            runs[first_idx].text = prefix + value + suffix
        else:
            runs[first_idx].text = prefix + value
            for i in covered[1:-1]:
                runs[i].text = ''
            runs[last_idx].text = suffix

        if key not in context:
            break
    return used


def replace_placeholders_in_doc(doc, context):
    used = set()
    for paragraph in _iter_all_paragraphs(doc):
        used |= replace_placeholders_in_paragraph(paragraph, context)
    return used


def _clear_header_content(header):
    hdr_element = header._element
    for child in list(hdr_element):
        if child.tag.endswith('}p') or child.tag.endswith('}tbl'):
            hdr_element.remove(child)
    return header.add_paragraph()


def apply_letterhead(doc, image_path):
    for section in doc.sections:
        section.header.is_linked_to_previous = False
        paragraph = _clear_header_content(section.header)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        available_width = section.page_width - section.left_margin - section.right_margin
        run = paragraph.add_run()
        run.add_picture(image_path, width=available_width)


def apply_typography(doc, font_family=None, font_size_pt=None, line_spacing=None):
    for paragraph in _iter_paragraphs(doc):
        if line_spacing is not None:
            paragraph.paragraph_format.line_spacing = line_spacing
        for run in paragraph.runs:
            if font_family:
                run.font.name = font_family
            if font_size_pt:
                run.font.size = Pt(font_size_pt)


def generate_document(template_path, output_path, context, letterhead_path=None,
                       font_family=None, font_size_pt=None, line_spacing=None):
    doc = Document(template_path)
    used = replace_placeholders_in_doc(doc, context)
    if letterhead_path:
        apply_letterhead(doc, letterhead_path)
    if font_family or font_size_pt or line_spacing:
        apply_typography(doc, font_family, font_size_pt, line_spacing)
    doc.save(output_path)
    return used


def format_value(value, tipo_campo):
    if value is None or value == '':
        return ''
    if tipo_campo == 'data':
        try:
            return datetime.strptime(str(value), '%Y-%m-%d').strftime('%d/%m/%Y')
        except ValueError:
            return str(value)
    if tipo_campo == 'importo':
        try:
            cents = int(value)
        except (TypeError, ValueError):
            return str(value)
        euros = cents / 100
        formatted = f'{euros:,.2f}'.replace(',', '#').replace('.', ',').replace('#', '.')
        return f'{formatted} €'
    if tipo_campo == 'numero':
        try:
            f = float(value)
        except (TypeError, ValueError):
            return str(value)
        if f == int(f):
            return str(int(f))
        return str(f).replace('.', ',')
    return str(value)


# ---------- endpoint ----------

def _handle(body):
    access_token = body.get('access_token')
    studio_id = get_studio_id(access_token)
    if not studio_id:
        raise PermissionError('Sessione non valida')

    matter_id = body.get('matter_id')
    template_id = body.get('template_id')
    manual_values = body.get('manual_values') or {}
    output_filename = body.get('output_filename') or 'documento.docx'
    if not matter_id or not template_id:
        raise ValueError('matter_id e template_id sono obbligatori')

    matter = _first(rest_get(f'/rest/v1/matters?id=eq.{matter_id}&studio_id=eq.{studio_id}&select=*'))
    if not matter:
        raise ValueError('Pratica non trovata')

    template = _first(rest_get(
        f'/rest/v1/templates?id=eq.{template_id}'
        f'&or=(studio_id.is.null,studio_id.eq.{studio_id})&select=*'
    ))
    if not template:
        raise ValueError('Template non trovato o non accessibile')

    client = _first(rest_get(f'/rest/v1/clients?id=eq.{matter["client_id"]}&select=*'))
    sinistro = None
    if matter.get('tipo_pratica') == 'sinistro':
        sinistro = _first(rest_get(f'/rest/v1/sinistri?matter_id=eq.{matter_id}&select=*'))

    placeholders = rest_get(f'/rest/v1/template_placeholders?template_id=eq.{template_id}&select=*&order=ordine')

    context = {}
    missing = []
    for p in placeholders:
        sorgente = p['sorgente']
        campo = p.get('campo_sorgente')
        key = p['placeholder_key']
        raw_value = None
        if sorgente == 'client' and campo:
            raw_value = (client or {}).get(campo)
        elif sorgente == 'matter' and campo:
            raw_value = matter.get(campo)
        elif sorgente == 'sinistro' and campo:
            raw_value = (sinistro or {}).get(campo)
        else:
            raw_value = manual_values.get(key)

        if raw_value in (None, '') and p.get('obbligatorio'):
            missing.append(p['etichetta'])
            continue
        context[key] = format_value(raw_value, p.get('tipo_campo'))

    if missing:
        raise ValueError('Campi obbligatori mancanti: ' + ', '.join(missing))

    template_scope = template['studio_id'] or SYSTEM_SCOPE
    template_blob = download_object(template['storage_path'])
    template_bytes = decrypt_bytes(template_blob, template_scope)

    settings = _first(rest_get(f'/rest/v1/studio_settings?studio_id=eq.{studio_id}&select=*')) or {}
    letterhead_path_tmp = None
    if settings.get('letterhead_storage_path'):
        letterhead_blob = download_object(settings['letterhead_storage_path'])
        letterhead_bytes = decrypt_bytes(letterhead_blob, studio_id)
        letterhead_path_tmp = os.path.join(tempfile.gettempdir(), f'letterhead_{uuid.uuid4().hex}.png')
        with open(letterhead_path_tmp, 'wb') as f:
            f.write(letterhead_bytes)

    template_path_tmp = os.path.join(tempfile.gettempdir(), f'template_{uuid.uuid4().hex}.docx')
    output_path_tmp = os.path.join(tempfile.gettempdir(), f'output_{uuid.uuid4().hex}.docx')
    with open(template_path_tmp, 'wb') as f:
        f.write(template_bytes)

    try:
        used = generate_document(
            template_path_tmp, output_path_tmp, context,
            letterhead_path=letterhead_path_tmp,
            font_family=settings.get('font_family') or 'Times New Roman',
            font_size_pt=float(settings.get('font_size_pt') or 12),
            line_spacing=float(settings.get('line_spacing') or 1.5),
        )
        with open(output_path_tmp, 'rb') as f:
            output_bytes = f.read()
    finally:
        for p in (template_path_tmp, output_path_tmp, letterhead_path_tmp):
            if p and os.path.isfile(p):
                os.remove(p)

    documento_id = str(uuid.uuid4())
    storage_path = f'documenti/{studio_id}/{documento_id}.docx.enc'
    upload_object(storage_path, encrypt_bytes(output_bytes, studio_id))

    rest_post('/rest/v1/documenti', {
        'id': documento_id, 'studio_id': studio_id, 'matter_id': matter_id, 'template_id': template_id,
        'nome_file': output_filename, 'storage_path': storage_path,
    })

    return {
        'ok': True,
        'documento_id': documento_id,
        'nome_file': output_filename,
        'placeholder_usati': sorted(used),
    }


class handler(BaseHTTPRequestHandler):
    def do_POST(self):
        try:
            length = int(self.headers.get('Content-Length', 0))
            body = json.loads(self.rfile.read(length).decode('utf-8') or '{}')
            result = _handle(body)
            self._send(200, result)
        except PermissionError as exc:
            self._send(401, {'ok': False, 'error': str(exc)})
        except ValueError as exc:
            self._send(400, {'ok': False, 'error': str(exc)})
        except Exception as exc:  # noqa: BLE001 - confine della funzione, non deve mai propagare
            self._send(500, {'ok': False, 'error': str(exc)})

    def _send(self, status, payload):
        self.send_response(status)
        self.send_header('Content-Type', 'application/json')
        self.end_headers()
        self.wfile.write(json.dumps(payload).encode('utf-8'))
